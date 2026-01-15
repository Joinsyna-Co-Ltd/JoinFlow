"""
Knowledge Base Management
========================

Document upload, indexing, and management for RAG system.
"""

import hashlib
import json
import logging
import mimetypes
import os
import re
import shutil
import uuid
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from enum import Enum

logger = logging.getLogger(__name__)


class DocumentStatus(str, Enum):
    PENDING = "pending"        # 待处理
    PROCESSING = "processing"  # 处理中
    INDEXED = "indexed"        # 已索引
    FAILED = "failed"          # 失败
    ARCHIVED = "archived"      # 已归档


class DocumentType(str, Enum):
    TEXT = "text"
    PDF = "pdf"
    MARKDOWN = "markdown"
    HTML = "html"
    CODE = "code"
    DOCX = "docx"
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class Document:
    """文档元数据"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    name: str = ""
    original_filename: str = ""
    
    # 分类
    collection: str = "default"  # 所属知识库
    tags: List[str] = field(default_factory=list)
    
    # 文件信息
    file_path: str = ""
    file_size: int = 0
    file_hash: str = ""
    mime_type: str = ""
    doc_type: DocumentType = DocumentType.UNKNOWN
    
    # 内容
    content: str = ""
    chunk_count: int = 0
    
    # 状态
    status: DocumentStatus = DocumentStatus.PENDING
    error_message: str = ""
    
    # 元数据
    created_at: datetime = field(default_factory=datetime.now)
    updated_at: datetime = field(default_factory=datetime.now)
    indexed_at: Optional[datetime] = None
    created_by: str = "default"
    
    # 额外信息
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> dict:
        data = asdict(self)
        data['status'] = self.status.value
        data['doc_type'] = self.doc_type.value
        data['created_at'] = self.created_at.isoformat()
        data['updated_at'] = self.updated_at.isoformat()
        data['indexed_at'] = self.indexed_at.isoformat() if self.indexed_at else None
        return data
    
    @classmethod
    def from_dict(cls, data: dict) -> "Document":
        if 'status' in data:
            data['status'] = DocumentStatus(data['status'])
        if 'doc_type' in data:
            data['doc_type'] = DocumentType(data['doc_type'])
        for field_name in ['created_at', 'updated_at', 'indexed_at']:
            if isinstance(data.get(field_name), str):
                data[field_name] = datetime.fromisoformat(data[field_name])
        return cls(**data)


@dataclass
class Collection:
    """知识库集合"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    name: str = ""
    description: str = ""
    
    # 配置
    chunk_size: int = 500        # 分块大小
    chunk_overlap: int = 50      # 重叠大小
    embedding_model: str = ""    # 嵌入模型
    
    # 统计
    document_count: int = 0
    chunk_count: int = 0
    total_size: int = 0
    
    # 元数据
    created_at: datetime = field(default_factory=datetime.now)
    updated_at: datetime = field(default_factory=datetime.now)
    created_by: str = "default"
    
    def to_dict(self) -> dict:
        data = asdict(self)
        data['created_at'] = self.created_at.isoformat()
        data['updated_at'] = self.updated_at.isoformat()
        return data
    
    @classmethod
    def from_dict(cls, data: dict) -> "Collection":
        for field_name in ['created_at', 'updated_at']:
            if isinstance(data.get(field_name), str):
                data[field_name] = datetime.fromisoformat(data[field_name])
        return cls(**data)


class TextChunker:
    """文本分块器"""
    
    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 500,
        overlap: int = 50
    ) -> List[Tuple[str, int, int]]:
        """
        将文本分成重叠的块
        
        Returns:
            List of (chunk_text, start_pos, end_pos)
        """
        if not text:
            return []
        
        chunks = []
        start = 0
        text_len = len(text)
        
        while start < text_len:
            end = min(start + chunk_size, text_len)
            
            # 尝试在句子边界分割
            if end < text_len:
                # 查找最近的句子结束符
                for sep in ['\n\n', '。', '！', '？', '.', '!', '?', '\n']:
                    pos = text.rfind(sep, start, end)
                    if pos > start + chunk_size // 2:
                        end = pos + len(sep)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append((chunk, start, end))
            
            start = end - overlap
        
        return chunks
    
    @staticmethod
    def chunk_by_paragraph(text: str, max_size: int = 1000) -> List[str]:
        """按段落分块"""
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(current) + len(para) + 2 <= max_size:
                current += ("\n\n" if current else "") + para
            else:
                if current:
                    chunks.append(current)
                current = para
        
        if current:
            chunks.append(current)
        
        return chunks


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def detect_type(filename: str, mime_type: str = "") -> DocumentType:
        """检测文档类型"""
        ext = Path(filename).suffix.lower()
        
        ext_map = {
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN,
            '.pdf': DocumentType.PDF,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.py': DocumentType.CODE,
            '.js': DocumentType.CODE,
            '.ts': DocumentType.CODE,
            '.java': DocumentType.CODE,
            '.cpp': DocumentType.CODE,
            '.c': DocumentType.CODE,
            '.go': DocumentType.CODE,
            '.rs': DocumentType.CODE,
            '.docx': DocumentType.DOCX,
            '.xlsx': DocumentType.XLSX,
            '.xls': DocumentType.XLSX,
            '.csv': DocumentType.CSV,
            '.json': DocumentType.JSON,
            '.png': DocumentType.IMAGE,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
            '.gif': DocumentType.IMAGE,
            '.webp': DocumentType.IMAGE,
        }
        
        return ext_map.get(ext, DocumentType.UNKNOWN)
    
    @staticmethod
    def parse_text(file_path: str) -> str:
        """解析纯文本"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def parse_markdown(file_path: str) -> str:
        """解析Markdown"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def parse_json(file_path: str) -> str:
        """解析JSON"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return json.dumps(data, ensure_ascii=False, indent=2)
    
    @staticmethod
    def parse_csv(file_path: str) -> str:
        """解析CSV"""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            return df.to_string()
        except ImportError:
            # 简单解析
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析PDF"""
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except ImportError:
            logger.warning("PyPDF2 not installed, cannot parse PDF")
            return ""
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            return ""
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析DOCX"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            logger.warning("python-docx not installed, cannot parse DOCX")
            return ""
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            return ""
    
    @classmethod
    def parse(cls, file_path: str, doc_type: DocumentType) -> str:
        """解析文档"""
        parsers = {
            DocumentType.TEXT: cls.parse_text,
            DocumentType.MARKDOWN: cls.parse_markdown,
            DocumentType.CODE: cls.parse_text,
            DocumentType.JSON: cls.parse_json,
            DocumentType.CSV: cls.parse_csv,
            DocumentType.PDF: cls.parse_pdf,
            DocumentType.DOCX: cls.parse_docx,
            DocumentType.HTML: cls.parse_text,
        }
        
        parser = parsers.get(doc_type)
        if parser:
            return parser(file_path)
        
        # 尝试作为文本解析
        try:
            return cls.parse_text(file_path)
        except Exception:
            return ""


class KnowledgeBaseManager:
    """知识库管理器"""
    
    def __init__(
        self,
        storage_path: str = "./knowledge_base",
        vector_store=None,
        embedder=None
    ):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)
        
        self.files_path = self.storage_path / "files"
        self.files_path.mkdir(exist_ok=True)
        
        self.vector_store = vector_store
        self.embedder = embedder
        
        self.collections: Dict[str, Collection] = {}
        self.documents: Dict[str, Document] = {}
        
        self._load_metadata()
    
    def _load_metadata(self):
        """加载元数据"""
        meta_file = self.storage_path / "metadata.json"
        if meta_file.exists():
            try:
                with open(meta_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                for col_data in data.get('collections', []):
                    col = Collection.from_dict(col_data)
                    self.collections[col.id] = col
                
                for doc_data in data.get('documents', []):
                    doc = Document.from_dict(doc_data)
                    self.documents[doc.id] = doc
                
                logger.info(f"Loaded {len(self.collections)} collections, {len(self.documents)} documents")
            except Exception as e:
                logger.error(f"Failed to load knowledge base metadata: {e}")
        
        # 确保默认集合存在
        if 'default' not in [c.name for c in self.collections.values()]:
            self.create_collection(Collection(name="default", description="默认知识库"))
    
    def _save_metadata(self):
        """保存元数据"""
        meta_file = self.storage_path / "metadata.json"
        try:
            with open(meta_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'collections': [c.to_dict() for c in self.collections.values()],
                    'documents': [d.to_dict() for d in self.documents.values()]
                }, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Failed to save knowledge base metadata: {e}")
    
    # ========== 集合管理 ==========
    
    def create_collection(self, collection: Collection) -> Collection:
        """创建知识库集合"""
        self.collections[collection.id] = collection
        self._save_metadata()
        logger.info(f"Created collection: {collection.name}")
        return collection
    
    def update_collection(self, collection_id: str, updates: dict) -> Optional[Collection]:
        """更新集合"""
        if collection_id not in self.collections:
            return None
        
        col = self.collections[collection_id]
        for key, value in updates.items():
            if hasattr(col, key):
                setattr(col, key, value)
        col.updated_at = datetime.now()
        
        self._save_metadata()
        return col
    
    def delete_collection(self, collection_id: str) -> bool:
        """删除集合"""
        if collection_id not in self.collections:
            return False
        
        # 删除集合中的所有文档
        docs_to_delete = [d.id for d in self.documents.values() 
                         if d.collection == self.collections[collection_id].name]
        for doc_id in docs_to_delete:
            self.delete_document(doc_id)
        
        del self.collections[collection_id]
        self._save_metadata()
        return True
    
    def get_collection(self, collection_id: str) -> Optional[Collection]:
        """获取集合"""
        return self.collections.get(collection_id)
    
    def get_collection_by_name(self, name: str) -> Optional[Collection]:
        """通过名称获取集合"""
        for col in self.collections.values():
            if col.name == name:
                return col
        return None
    
    def list_collections(self) -> List[Collection]:
        """列出所有集合"""
        return list(self.collections.values())
    
    # ========== 文档管理 ==========
    
    def add_document(
        self,
        file_content: bytes,
        filename: str,
        collection: str = "default",
        tags: List[str] = None,
        metadata: Dict = None,
        user_id: str = "default"
    ) -> Document:
        """添加文档"""
        # 计算文件hash
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # 检查重复
        for doc in self.documents.values():
            if doc.file_hash == file_hash and doc.collection == collection:
                logger.info(f"Document already exists: {filename}")
                return doc
        
        # 创建文档记录
        doc = Document(
            name=Path(filename).stem,
            original_filename=filename,
            collection=collection,
            tags=tags or [],
            file_size=len(file_content),
            file_hash=file_hash,
            mime_type=mimetypes.guess_type(filename)[0] or "",
            doc_type=DocumentParser.detect_type(filename),
            status=DocumentStatus.PENDING,
            created_by=user_id,
            metadata=metadata or {}
        )
        
        # 保存文件
        file_path = self.files_path / f"{doc.id}{Path(filename).suffix}"
        with open(file_path, 'wb') as f:
            f.write(file_content)
        doc.file_path = str(file_path)
        
        self.documents[doc.id] = doc
        self._save_metadata()
        
        # 更新集合统计
        col = self.get_collection_by_name(collection)
        if col:
            col.document_count += 1
            col.total_size += len(file_content)
            col.updated_at = datetime.now()
            self._save_metadata()
        
        logger.info(f"Added document: {filename} to {collection}")
        return doc
    
    def add_text_document(
        self,
        content: str,
        name: str,
        collection: str = "default",
        tags: List[str] = None,
        user_id: str = "default"
    ) -> Document:
        """添加纯文本文档"""
        return self.add_document(
            file_content=content.encode('utf-8'),
            filename=f"{name}.txt",
            collection=collection,
            tags=tags,
            user_id=user_id
        )
    
    def process_document(self, doc_id: str) -> bool:
        """处理文档（解析和索引）"""
        doc = self.documents.get(doc_id)
        if not doc:
            return False
        
        try:
            doc.status = DocumentStatus.PROCESSING
            self._save_metadata()
            
            # 解析文档
            content = DocumentParser.parse(doc.file_path, doc.doc_type)
            doc.content = content
            
            # 分块
            col = self.get_collection_by_name(doc.collection)
            chunk_size = col.chunk_size if col else 500
            chunk_overlap = col.chunk_overlap if col else 50
            
            chunks = TextChunker.chunk_text(content, chunk_size, chunk_overlap)
            doc.chunk_count = len(chunks)
            
            # 索引到向量数据库
            if self.vector_store and self.embedder and chunks:
                ids = [f"{doc.id}_{i}" for i in range(len(chunks))]
                vectors = [self.embedder.embed(chunk[0]) for chunk in chunks]
                metadata_list = [
                    {
                        "doc_id": doc.id,
                        "doc_name": doc.name,
                        "collection": doc.collection,
                        "chunk_index": i,
                        "start_pos": chunk[1],
                        "end_pos": chunk[2],
                        "content": chunk[0][:500]  # 存储部分内容
                    }
                    for i, chunk in enumerate(chunks)
                ]
                
                self.vector_store.upsert(ids, vectors, metadata_list)
                
                # 更新集合统计
                if col:
                    col.chunk_count += len(chunks)
                    self._save_metadata()
            
            doc.status = DocumentStatus.INDEXED
            doc.indexed_at = datetime.now()
            doc.updated_at = datetime.now()
            self._save_metadata()
            
            logger.info(f"Processed document: {doc.name} ({len(chunks)} chunks)")
            return True
            
        except Exception as e:
            logger.error(f"Failed to process document {doc_id}: {e}")
            doc.status = DocumentStatus.FAILED
            doc.error_message = str(e)
            self._save_metadata()
            return False
    
    def delete_document(self, doc_id: str) -> bool:
        """删除文档"""
        doc = self.documents.get(doc_id)
        if not doc:
            return False
        
        # 删除文件
        if doc.file_path and Path(doc.file_path).exists():
            Path(doc.file_path).unlink()
        
        # 从向量数据库删除
        if self.vector_store:
            # TODO: 实现向量数据库删除
            pass
        
        # 更新集合统计
        col = self.get_collection_by_name(doc.collection)
        if col:
            col.document_count = max(0, col.document_count - 1)
            col.total_size = max(0, col.total_size - doc.file_size)
            col.chunk_count = max(0, col.chunk_count - doc.chunk_count)
        
        del self.documents[doc_id]
        self._save_metadata()
        
        logger.info(f"Deleted document: {doc.name}")
        return True
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """获取文档"""
        return self.documents.get(doc_id)
    
    def list_documents(
        self,
        collection: str = None,
        status: DocumentStatus = None,
        tags: List[str] = None,
        limit: int = 100,
        offset: int = 0
    ) -> Tuple[List[Document], int]:
        """列出文档"""
        docs = list(self.documents.values())
        
        if collection:
            docs = [d for d in docs if d.collection == collection]
        if status:
            docs = [d for d in docs if d.status == status]
        if tags:
            docs = [d for d in docs if any(t in d.tags for t in tags)]
        
        total = len(docs)
        docs = sorted(docs, key=lambda d: d.created_at, reverse=True)
        docs = docs[offset:offset + limit]
        
        return docs, total
    
    def search_documents(
        self,
        query: str,
        collection: str = None,
        top_k: int = 10
    ) -> List[Dict]:
        """搜索文档"""
        if not self.vector_store or not self.embedder:
            return []
        
        query_vector = self.embedder.embed(query)
        
        # 构建过滤器
        filters = None
        if collection:
            from qdrant_client.models import Filter, FieldCondition, MatchValue
            filters = Filter(must=[
                FieldCondition(key="collection", match=MatchValue(value=collection))
            ])
        
        results = self.vector_store.search(query_vector, top_k, filters)
        
        return [
            {
                "doc_id": r.payload.get("doc_id"),
                "doc_name": r.payload.get("doc_name"),
                "collection": r.payload.get("collection"),
                "content": r.payload.get("content"),
                "score": r.score
            }
            for r in results
        ]
    
    def get_statistics(self) -> Dict:
        """获取统计信息"""
        return {
            "total_collections": len(self.collections),
            "total_documents": len(self.documents),
            "total_size": sum(d.file_size for d in self.documents.values()),
            "total_chunks": sum(d.chunk_count for d in self.documents.values()),
            "by_status": {
                status.value: len([d for d in self.documents.values() if d.status == status])
                for status in DocumentStatus
            },
            "by_type": {
                doc_type.value: len([d for d in self.documents.values() if d.doc_type == doc_type])
                for doc_type in DocumentType
                if any(d.doc_type == doc_type for d in self.documents.values())
            }
        }

